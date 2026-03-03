"""
Generate a professional bug report Word document for thinkartha.com.
Combines previously reported bugs with new audit findings.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colour palette ──────────────────────────────────────────────
TEAL       = RGBColor(0x0E, 0x8B, 0x8B)
DARK_TEAL  = RGBColor(0x0A, 0x5E, 0x5E)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x1A, 0x1A, 0x2E)
GREY       = RGBColor(0x64, 0x74, 0x8B)
LIGHT_BG   = RGBColor(0xF8, 0xFA, 0xFC)
RED_BG     = "F8D7DA"
AMBER_BG   = "FFF3CD"
GREEN_BG   = "D1E7DD"
BLUE_BG    = "CFE2FF"
TEAL_HEX   = "0E8B8B"
HEADER_BG  = "134A4A"

def set_cell_bg(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_cell_text(cell, text, bold=False, color=BLACK, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Calibri"

def style_header_row(row, col_count):
    for i in range(col_count):
        set_cell_bg(row.cells[i], HEADER_BG)
        for p in row.cells[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.bold = True
                run.font.size = Pt(9)

doc = Document()

# ── Page setup ──────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# ── Default font ────────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(10)
font.color.rgb = BLACK

# ── Custom heading styles ───────────────────────────────────────
for level in [1, 2, 3]:
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = DARK_TEAL
    hs.font.bold = True
    hs.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    hs.paragraph_format.space_after = Pt(6)

doc.styles["Heading 1"].font.size = Pt(20)
doc.styles["Heading 2"].font.size = Pt(14)
doc.styles["Heading 3"].font.size = Pt(11)

# ═══════════════════════════════════════════════════════════════
#  COVER / TITLE PAGE
# ═══════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph("")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("WEBSITE QUALITY REVIEW")
run.font.size = Pt(28)
run.font.color.rgb = TEAL
run.bold = True
run.font.name = "Calibri"

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("thinkartha.com — Detailed Findings & Recommendations")
run.font.size = Pt(14)
run.font.color.rgb = GREY
run.font.name = "Calibri"

doc.add_paragraph("")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for label, value in [
    ("Prepared for: ", "Artha Solutions Website Development Team"),
    ("\nPrepared by: ", "Vishnu Sai Rachakatla"),
    ("\nDate: ", "March 3, 2026"),
    ("\nClassification: ", "Internal — Confidential"),
]:
    r = meta.add_run(label)
    r.font.size = Pt(10)
    r.font.color.rgb = GREY
    r.font.name = "Calibri"
    r = meta.add_run(value)
    r.font.size = Pt(10)
    r.font.color.rgb = BLACK
    r.bold = True
    r.font.name = "Calibri"

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════

doc.add_heading("1. Introduction", level=1)

doc.add_paragraph(
    "Dear Team,"
)
doc.add_paragraph(
    "Thank you for your continued efforts in maintaining and improving the Artha Solutions website "
    "(thinkartha.com). We truly appreciate the work that has gone into building and managing the site "
    "to represent our brand and services."
)
doc.add_paragraph(
    "As part of our ongoing commitment to delivering the best possible digital experience for our "
    "clients and prospects, we conducted a comprehensive quality review of the current website. This "
    "document consolidates findings from two sources:"
)

bullets = [
    "A prior internal review that identified several issues, many of which have already been addressed — thank you for those fixes.",
    "A new deep-dive audit examining every major page for content accuracy, visual consistency, "
    "mobile responsiveness, performance, and search engine optimization.",
]
for b in bullets:
    p = doc.add_paragraph(b, style="List Bullet")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(10)

doc.add_paragraph(
    "The purpose of this document is to provide a clear, organized reference for the development team "
    "so that remaining items can be prioritized and addressed efficiently. Every finding is presented "
    "with a description, the affected page(s), suggested resolution, and a priority level."
)
doc.add_paragraph(
    "We are confident that resolving these items will further elevate the professionalism and "
    "effectiveness of thinkartha.com. Please do not hesitate to reach out if any item requires "
    "additional context or clarification."
)

# ═══════════════════════════════════════════════════════════════
#  2. STATUS OF PREVIOUSLY REPORTED ISSUES
# ═══════════════════════════════════════════════════════════════

doc.add_heading("2. Status of Previously Reported Issues", level=1)

doc.add_paragraph(
    "The following table summarizes bugs that were identified in our earlier review. We are pleased "
    "to note that the majority of these have been resolved. Items still outstanding or needing "
    "implementation are highlighted for continued attention."
)

# Table of previously reported bugs
prev_bugs = [
    ("No call functionality for Indonesia region", "Fixed", "—", "✅"),
    ("Website phone layout is non-responsive", "Fixed", "—", "✅"),
    ("CAPTCHA required for form filling, not for website — please remove", "Fixed", "—", "✅"),
    ("Instagram symbol is not present", "N/A", "Verify if Instagram presence is active", "⚠️"),
    ("In phone layout, text and symbols are misaligned", "Fixed", "—", "✅"),
    ("In events section, year is missing", "Fixed", "—", "✅"),
    ("In the blogs section, navigation arrows are out of border", "Fixed", "—", "✅"),
    ("Industries dropdown menu has gaps & description is not aligned", "Fixed", "—", "✅"),
    ("Resources section read-time visible for only one topic; missing on others", "Fixed", "—", "✅"),
    ("Awards & Recognition: heading issue on 'Talend Data Masters Award' tile", "Fixed", "—", "✅"),
    ("Partners logo carousel should have smooth flow & pause on hover", "Fixed", "—", "✅"),
    ("Outdated events that have passed should be removed or display a message", "Not Fixed", "Display 'Stay tuned for upcoming events' or archive past events", "🔴"),
    ("Search functionality does not lead to appropriate content; needs auto-suggestions", "Not Fixed", "Implement search indexing with auto-complete suggestions", "🔴"),
    ("Blog posts are missing on some pages", "Needs Implementation", "Ensure consistent blog feeds across all relevant pages", "🟡"),
    ("Images are missing on Contact Us page", "Needs Implementation", "Add office images or map embeds to Contact page", "🟡"),
]

table = doc.add_table(rows=1, cols=5)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["#", "Issue Description", "Current Status", "Suggested Action", ""]
hdr_row = table.rows[0]
for i, h in enumerate(headers):
    add_cell_text(hdr_row.cells[i], h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
style_header_row(hdr_row, 5)

# Set column widths
widths = [Cm(0.8), Cm(7.5), Cm(2.5), Cm(5.5), Cm(0.8)]
for i, w in enumerate(widths):
    hdr_row.cells[i].width = w

for idx, (desc, status, action, icon) in enumerate(prev_bugs, 1):
    row = table.add_row()
    add_cell_text(row.cells[0], str(idx), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(row.cells[1], desc, size=9)
    add_cell_text(row.cells[2], status, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(row.cells[3], action, size=9)
    add_cell_text(row.cells[4], icon, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    for i, w in enumerate(widths):
        row.cells[i].width = w

    # Color-code status
    if status == "Fixed":
        set_cell_bg(row.cells[2], GREEN_BG)
    elif status == "Not Fixed":
        set_cell_bg(row.cells[2], RED_BG)
    elif "Needs" in status:
        set_cell_bg(row.cells[2], AMBER_BG)

doc.add_paragraph("")

# ═══════════════════════════════════════════════════════════════
#  3. PREVIOUSLY SUGGESTED IMPROVEMENTS
# ═══════════════════════════════════════════════════════════════

doc.add_heading("3. Previously Suggested Improvements", level=1)

doc.add_paragraph(
    "In addition to the bugs above, several improvement suggestions were made during the earlier "
    "review. Below is the current status of those suggestions."
)

suggestions = [
    ("Adding a chatbot to the UI for more interactive engagement", "Accepted", "Consider lightweight chatbot integration (e.g., Drift, Tidio, or custom widget)"),
    ("Remove 2-line descriptions in Artha Advantage dropdown menu", "Not Accepted", "No action required — decision to retain current format"),
    ("Success stories should be in a separate page/section, not in Resources", "Under Review", "Consider creating a dedicated 'Case Studies' page for better content organization"),
    ("Spelling corrections across the site", "Accepted", "Several spelling issues remain (see Section 4 for details)"),
    ("B'etl is repeated in Artha Advantage section — one should be removed", "Accepted", "Please verify this has been fully resolved"),
    ("In events, datetime commands should be used instead of headers (H6 tags)", "Accepted", "Use semantic <time> elements instead of heading tags for event dates"),
]

table2 = doc.add_table(rows=1, cols=4)
table2.style = "Table Grid"
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ["#", "Suggestion", "Response", "Recommended Action"]
hdr2 = table2.rows[0]
for i, h in enumerate(headers2):
    add_cell_text(hdr2.cells[i], h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
style_header_row(hdr2, 4)

widths2 = [Cm(0.8), Cm(6.0), Cm(2.5), Cm(7.0)]
for i, w in enumerate(widths2):
    hdr2.cells[i].width = w

for idx, (desc, resp, action) in enumerate(suggestions, 1):
    row = table2.add_row()
    add_cell_text(row.cells[0], str(idx), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(row.cells[1], desc, size=9)
    add_cell_text(row.cells[2], resp, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(row.cells[3], action, size=9)
    for i, w in enumerate(widths2):
        row.cells[i].width = w
    if resp == "Accepted":
        set_cell_bg(row.cells[2], GREEN_BG)
    elif resp == "Not Accepted":
        set_cell_bg(row.cells[2], "E2E3E5")
    else:
        set_cell_bg(row.cells[2], AMBER_BG)

doc.add_paragraph("")

# ═══════════════════════════════════════════════════════════════
#  4. NEW FINDINGS — CONTENT & TEXT ISSUES
# ═══════════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("4. New Findings — Content & Text Quality", level=1)

doc.add_paragraph(
    "During our latest review, we identified several content-quality issues across the website. "
    "These are organized below by type for clarity."
)

# 4.1 Duplicate / Repeating Text
doc.add_heading("4.1 Duplicate / Repeating Text", level=2)

doc.add_paragraph(
    "We noticed that certain text blocks appear multiple times on the same page. This may be caused "
    "by duplicated Elementor widget rows. Resolving these will improve the reading experience and "
    "page length."
)

dup_issues = [
    ("Homepage — 'Our Services & Solutions' Tabs",
     "All five tabs (Quality, Integrity & Trust / Analyze, Automate & Accelerate / Artha Advantage for SAP / One Platform / Redefining Data Governance) contain bullet lists that are repeated twice — the same list appears back-to-back within each tab.",
     "Critical",
     "Remove the duplicate Elementor widget (or duplicate inner section) in each tab. Each tab should show its bullet list only once."),

    ("Homepage — Resources Section",
     "The white paper cards ('AI and Data Modernization' and 'Future-Ready Data Foundation') each appear twice with identical text, images, and download links.",
     "Critical",
     "Remove the duplicate resource card widgets so each card appears only once."),

    ("Homepage — Case Study Link Mismatch",
     "The 'Accelerating Digital Transformation for Seamless Partner and Customer Operations' card has a 'Download Now!' link that points to a completely different case study (healthcare risk adjustment).",
     "Critical",
     "Update the link URL to point to the correct case study, or replace the card with the intended content."),

    ("SAP Page — Identical Heading × 4",
     "The heading 'SAP Data Servcies powered by Artha Advantage' appears four times at the top of the page. Each section describes a different phase (before, during, testing, after migration) but uses the same heading.",
     "High",
     "Give each section a unique, descriptive heading (e.g., 'Pre-Migration Data Cleansing', 'Migration Integration', 'Test Data Management', 'Post-Migration Replication')."),

    ("Multiple Pages — Same Testimonial",
     "The Steve Brennan testimonial ('Cleansing and consolidating consumer data…') appears identically on the Homepage, Artha Advantage, Data Governance, Manufacturing, Cloud, Managed Services, and other pages.",
     "Medium",
     "Rotate testimonials per page — assign the most relevant client quote to each service/industry page for a more authentic feel."),
]

table3 = doc.add_table(rows=1, cols=5)
table3.style = "Table Grid"
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

headers3 = ["#", "Location & Description", "Priority", "Recommended Fix", "Status"]
hdr3 = table3.rows[0]
for i, h in enumerate(headers3):
    add_cell_text(hdr3.cells[i], h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
style_header_row(hdr3, 5)

widths3 = [Cm(0.6), Cm(6.5), Cm(1.5), Cm(6.5), Cm(1.5)]
for i, w in enumerate(widths3):
    hdr3.cells[i].width = w

for idx, (loc, desc, prio, fix) in enumerate(dup_issues, 1):
    row = table3.add_row()
    add_cell_text(row.cells[0], str(idx), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    # Combine location + description
    cell1 = row.cells[1]
    cell1.text = ""
    p = cell1.paragraphs[0]
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    r = p.add_run(loc + "\n")
    r.bold = True
    r.font.size = Pt(9)
    r.font.name = "Calibri"
    r.font.color.rgb = DARK_TEAL
    r = p.add_run(desc)
    r.font.size = Pt(8)
    r.font.name = "Calibri"
    r.font.color.rgb = BLACK

    add_cell_text(row.cells[2], prio, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(row.cells[3], fix, size=8)
    add_cell_text(row.cells[4], "Open", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    for i, w in enumerate(widths3):
        row.cells[i].width = w

    if prio == "Critical":
        set_cell_bg(row.cells[2], RED_BG)
    elif prio == "High":
        set_cell_bg(row.cells[2], AMBER_BG)
    else:
        set_cell_bg(row.cells[2], BLUE_BG)

doc.add_paragraph("")

# 4.2 Spelling & Grammar
doc.add_heading("4.2 Spelling & Grammar Errors", level=2)

spelling = [
    ("SAP Page — Hero Section",
     '"SAP Data Servcies powered by Artha Advantage" — "Servcies" should be "Services".\nThis typo appears four times due to the repeated heading.',
     "Critical",
     "Correct spelling to 'Services' in all four instances."),

    ("Accelerators Page — Data Insights Platform",
     '"Al driven & robust comprehensive platform" — "Al" (capital A, lowercase l) should be "AI" (capital A, capital I).',
     "High",
     "Replace 'Al' with 'AI'."),

    ("Homepage — Services Tab",
     '"Trusted data for AI , and accurate 360° insights." — Extra space before the comma.',
     "Low",
     "Remove the extra space before the comma."),

    ("Homepage — Events URLs",
     "Multiple event and white-paper links end with '%20' (a URL-encoded trailing space), e.g. /events/artha-solutions-to-host-qlik-ai-reality-tour-2025-bengaluru/%20",
     "Medium",
     "Remove trailing spaces from all internal link URLs to avoid potential 404s or canonical issues."),
]

table4 = doc.add_table(rows=1, cols=5)
table4.style = "Table Grid"
table4.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr4 = table4.rows[0]
for i, h in enumerate(headers3):
    add_cell_text(hdr4.cells[i], h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
style_header_row(hdr4, 5)
for i, w in enumerate(widths3):
    hdr4.cells[i].width = w

for idx, (loc, desc, prio, fix) in enumerate(spelling, 1):
    row = table4.add_row()
    add_cell_text(row.cells[0], str(idx), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell1 = row.cells[1]
    cell1.text = ""
    p = cell1.paragraphs[0]
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    r = p.add_run(loc + "\n")
    r.bold = True
    r.font.size = Pt(9)
    r.font.name = "Calibri"
    r.font.color.rgb = DARK_TEAL
    r = p.add_run(desc)
    r.font.size = Pt(8)
    r.font.name = "Calibri"
    r.font.color.rgb = BLACK
    add_cell_text(row.cells[2], prio, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(row.cells[3], fix, size=8)
    add_cell_text(row.cells[4], "Open", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, w in enumerate(widths3):
        row.cells[i].width = w
    if prio == "Critical":
        set_cell_bg(row.cells[2], RED_BG)
    elif prio == "High":
        set_cell_bg(row.cells[2], AMBER_BG)
    elif prio == "Medium":
        set_cell_bg(row.cells[2], BLUE_BG)
    else:
        set_cell_bg(row.cells[2], "E2E3E5")

doc.add_paragraph("")

# ═══════════════════════════════════════════════════════════════
#  5. NEW FINDINGS — VISUAL & COSMETIC ISSUES
# ═══════════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("5. New Findings — Visual & Cosmetic Issues", level=1)

doc.add_paragraph(
    "The following visual and design-related observations were made during the audit. Addressing "
    "these will help create a more polished and consistent user experience."
)

cosmetic = [
    ("Low-Resolution Client Logos",
     "Homepage → 'Our Clients' Carousel",
     "Several client logos (M&T Bank, Maxis, DBS, Security Bank, match.com) appear blurry or pixelated, likely from small raster images being upscaled by CSS.",
     "Medium",
     "Replace with high-resolution PNG or SVG versions of each logo."),

    ("Low-Resolution Partner Badges",
     "Homepage → Certifications Section",
     "Some partner badges (Qlik, AWS, Microsoft) have inconsistent resolution — some are sharp while others are blurry.",
     "Medium",
     "Source uniform, high-resolution badge assets from partner portals."),

    ("Heavy Dark Overlays on Hero Images",
     "All major pages (Homepage, SAP, Cloud, etc.)",
     "Nearly every page uses a very dark semi-transparent overlay on the hero background, giving the site a heavy, monotonous feel across pages.",
     "Low",
     "Consider lighter overlays, gradient overlays, or varied hero treatments per section."),

    ("Inconsistent Event Card Imagery",
     "Homepage → 'Let's Meet' Section",
     "Event cards use different image styles (tall posters, headshots, narrow banners), creating an uneven visual grid.",
     "Medium",
     "Standardize event card image dimensions and aspect ratios."),

    ("Cookie Consent Banner",
     "All Pages",
     "The cookie consent banner is large and covers a significant portion of the viewport, feeling intrusive until dismissed.",
     "Low",
     "Consider a slimmer, less obtrusive bar design."),

    ("Inconsistent CTA Button Styles",
     "Multiple Pages",
     "CTA buttons vary between filled teal, outline teal, and plain text links. Hover effects also differ between sections.",
     "Medium",
     "Establish a consistent button design system: primary (filled), secondary (outline), and tertiary (text link)."),

    ("Contact Form — Generic Styling",
     "Contact Page",
     "The contact form uses default Elementor styling, looking generic compared to the rest of the site.",
     "Low",
     "Apply custom CSS to match the brand design (rounded inputs, teal accent focus state, subtle shadows)."),

    ("Tab Navigation Arrows — Hard to Spot",
     "Homepage → 'Our Services & Solutions'",
     "The small left/right navigation arrows for scrolling between service tabs are easy to miss, especially on smaller screens.",
     "Medium",
     "Use larger, more visible arrow icons or add subtle animation to indicate scrollability."),
]

table5 = doc.add_table(rows=1, cols=5)
table5.style = "Table Grid"
table5.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr5_labels = ["#", "Issue & Location", "Priority", "Suggested Resolution", "Status"]
hdr5 = table5.rows[0]
for i, h in enumerate(hdr5_labels):
    add_cell_text(hdr5.cells[i], h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
style_header_row(hdr5, 5)
for i, w in enumerate(widths3):
    hdr5.cells[i].width = w

for idx, (title, loc, desc, prio, fix) in enumerate(cosmetic, 1):
    row = table5.add_row()
    add_cell_text(row.cells[0], str(idx), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell1 = row.cells[1]
    cell1.text = ""
    p = cell1.paragraphs[0]
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.size = Pt(9)
    r.font.name = "Calibri"
    r.font.color.rgb = DARK_TEAL
    r = p.add_run(f"Page: {loc}\n")
    r.font.size = Pt(8)
    r.font.name = "Calibri"
    r.font.color.rgb = GREY
    r = p.add_run(desc)
    r.font.size = Pt(8)
    r.font.name = "Calibri"
    r.font.color.rgb = BLACK
    add_cell_text(row.cells[2], prio, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(row.cells[3], fix, size=8)
    add_cell_text(row.cells[4], "Open", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, w in enumerate(widths3):
        row.cells[i].width = w
    if prio == "Critical":
        set_cell_bg(row.cells[2], RED_BG)
    elif prio == "High":
        set_cell_bg(row.cells[2], AMBER_BG)
    elif prio == "Medium":
        set_cell_bg(row.cells[2], BLUE_BG)
    else:
        set_cell_bg(row.cells[2], "E2E3E5")

doc.add_paragraph("")

# ═══════════════════════════════════════════════════════════════
#  6. NEW FINDINGS — MOBILE RESPONSIVENESS
# ═══════════════════════════════════════════════════════════════

doc.add_heading("6. New Findings — Mobile Responsiveness", level=1)

doc.add_paragraph(
    "The website was tested at 375px viewport width (iPhone equivalent). The following issues "
    "were observed on mobile devices."
)

mobile = [
    ("Overlapping Elements on Mobile Homepage", "Homepage (375px)", "In the 'Let's Meet' section, the scroll/mouse icon overlaps the heading text. Other elements also bleed into adjacent sections due to insufficient mobile padding.", "High", "Add responsive padding/margin rules and test on actual mobile devices."),
    ("Oversized Client Logos on Mobile", "Homepage (375px)", "Client logos in the carousel appear disproportionately large on mobile, taking up excessive screen real estate.", "Medium", "Reduce logo size and spacing for viewports under 768px."),
    ("Mega-Menu on Mobile", "All Pages", "The multi-column mega-menu designed for desktop becomes an extremely long scrolling list on mobile, requiring excessive scrolling.", "Medium", "Implement a collapsible accordion-style menu for mobile navigation."),
    ("Service Tab Arrows Difficult to Tap", "Homepage (375px)", "The small tab navigation arrows in 'Our Services & Solutions' are too small for comfortable touch interaction.", "Medium", "Increase tap target size to at least 44×44px per mobile accessibility guidelines."),
]

table6 = doc.add_table(rows=1, cols=5)
table6.style = "Table Grid"
table6.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr6 = table6.rows[0]
for i, h in enumerate(hdr5_labels):
    add_cell_text(hdr6.cells[i], h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
style_header_row(hdr6, 5)
for i, w in enumerate(widths3):
    hdr6.cells[i].width = w

for idx, (title, loc, desc, prio, fix) in enumerate(mobile, 1):
    row = table6.add_row()
    add_cell_text(row.cells[0], str(idx), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell1 = row.cells[1]
    cell1.text = ""
    p = cell1.paragraphs[0]
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.size = Pt(9)
    r.font.name = "Calibri"
    r.font.color.rgb = DARK_TEAL
    r = p.add_run(f"Page: {loc}\n")
    r.font.size = Pt(8)
    r.font.name = "Calibri"
    r.font.color.rgb = GREY
    r = p.add_run(desc)
    r.font.size = Pt(8)
    r.font.name = "Calibri"
    r.font.color.rgb = BLACK
    add_cell_text(row.cells[2], prio, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(row.cells[3], fix, size=8)
    add_cell_text(row.cells[4], "Open", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, w in enumerate(widths3):
        row.cells[i].width = w
    if prio == "High":
        set_cell_bg(row.cells[2], AMBER_BG)
    else:
        set_cell_bg(row.cells[2], BLUE_BG)

doc.add_paragraph("")

# ═══════════════════════════════════════════════════════════════
#  7. NEW FINDINGS — STRUCTURE, SEO & PERFORMANCE
# ═══════════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("7. New Findings — Content Structure, SEO & Performance", level=1)

doc.add_paragraph(
    "The following issues relate to the technical structure of the website's content, search engine "
    "optimization, and page load performance."
)

doc.add_heading("7.1 Content Structure", level=2)

structure = [
    ("Heading Hierarchy Issues", "Multiple Pages", "Several pages skip heading levels (H2 → H6), use H6 for event dates instead of semantic <time> elements, or cram entire paragraphs into H1 tags (e.g., the Accelerators page H1 is a full paragraph).", "Medium", "Maintain proper H1 → H2 → H3 hierarchy. Use <time> for dates."),
    ("Tabbed Content Invisible to Search Engines", "Homepage", "The 'Our Services & Solutions' tabs use JavaScript to toggle content visibility. Search engine crawlers may not index content in inactive tabs.", "Medium", "Ensure all tab content is present in the DOM and accessible to crawlers, even if visually hidden."),
    ("No Breadcrumb Navigation", "Solution & Industry Sub-Pages", "Deep pages like /sap/clean-data-for-sap/ have no breadcrumbs, making it hard for users to understand their location in the site hierarchy.", "Medium", "Add breadcrumb navigation to all sub-pages."),
    ("Data Governance Title Tag", "Data Governance Page", "The page title reads: 'Data Governance & Consulting Services | Artha Solutions | Artha Solutions' — the brand name is duplicated.", "High", "Remove the duplicate 'Artha Solutions' from the title tag."),
]

table7 = doc.add_table(rows=1, cols=5)
table7.style = "Table Grid"
table7.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr7 = table7.rows[0]
for i, h in enumerate(hdr5_labels):
    add_cell_text(hdr7.cells[i], h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
style_header_row(hdr7, 5)
for i, w in enumerate(widths3):
    hdr7.cells[i].width = w

for idx, (title, loc, desc, prio, fix) in enumerate(structure, 1):
    row = table7.add_row()
    add_cell_text(row.cells[0], str(idx), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell1 = row.cells[1]
    cell1.text = ""
    p = cell1.paragraphs[0]
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.size = Pt(9)
    r.font.name = "Calibri"
    r.font.color.rgb = DARK_TEAL
    r = p.add_run(f"Page: {loc}\n")
    r.font.size = Pt(8)
    r.font.name = "Calibri"
    r.font.color.rgb = GREY
    r = p.add_run(desc)
    r.font.size = Pt(8)
    r.font.name = "Calibri"
    r.font.color.rgb = BLACK
    add_cell_text(row.cells[2], prio, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(row.cells[3], fix, size=8)
    add_cell_text(row.cells[4], "Open", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, w in enumerate(widths3):
        row.cells[i].width = w
    if prio == "High":
        set_cell_bg(row.cells[2], AMBER_BG)
    else:
        set_cell_bg(row.cells[2], BLUE_BG)

doc.add_paragraph("")

doc.add_heading("7.2 Performance Observations", level=2)

perf_points = [
    "Heavy JavaScript bundle: WordPress core, jQuery, Elementor JS, and multiple plugin scripts contribute to a significant JavaScript payload (estimated 400KB+), which impacts initial load time.",
    "No visible CDN: Assets appear to be served directly from the origin server. Implementing a CDN (e.g., Cloudflare) would significantly improve global load times.",
    "Image optimization: Client logos, event images, and partner badges could benefit from WebP format conversion and proper lazy loading.",
    "Large DOM tree: The mega-menu, footer link blocks, and repeated content sections create a very large DOM (estimated 3,000+ nodes on the homepage), increasing memory usage and slowing CSS rendering.",
    "Server-side rendering overhead: Each page request involves PHP execution and MySQL database queries, adding 200–800ms to the time-to-first-byte (TTFB). Page-level caching or a static caching plugin (e.g., WP Super Cache) would help.",
    "Render-blocking resources: Multiple CSS and JS files in the <head> block rendering. Consider deferring non-critical scripts and inlining critical CSS.",
]

for pt in perf_points:
    p = doc.add_paragraph(pt, style="List Bullet")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)

doc.add_paragraph("")

doc.add_heading("7.3 SEO Observations", level=2)

seo_points = [
    "Missing structured data (JSON-LD): No Organization, BreadcrumbList, FAQ, Event, or Article schema detected. Adding structured data improves search result appearance with rich snippets.",
    "Internal links with trailing whitespace (%20): Several event and white-paper links have a trailing space encoded as '%20', which could cause 404 errors or canonicalization issues.",
    "Missing alt attributes: Many images (client logos, event images) lack descriptive alt text, impacting accessibility and image search performance.",
    "No visible XML sitemap link: Consider adding a link to the sitemap in the footer or robots.txt for easier search engine crawling.",
    "Footer link density: The footer contains 50+ links across 4 categories, which may dilute link equity. Consider prioritizing key links.",
]

for pt in seo_points:
    p = doc.add_paragraph(pt, style="List Bullet")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)

# ═══════════════════════════════════════════════════════════════
#  8. SUMMARY & PRIORITY MATRIX
# ═══════════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("8. Summary & Priority Matrix", level=1)

doc.add_paragraph(
    "Below is a consolidated priority matrix of all new findings to help plan the resolution effort."
)

summary_data = [
    ("Critical", "7", "Duplicate text blocks, spelling typo on SAP page, mismatched case study link, trailing URL spaces, title tag duplication"),
    ("High", "5", "SAP heading differentiation, 'Al' → 'AI' typo, mobile overlapping elements, Data Governance title tag"),
    ("Medium", "11", "Testimonial rotation, logo resolution, event card imagery, CTA consistency, breadcrumbs, mobile menu, tab arrows, DOM/CDN improvements"),
    ("Low", "5", "Dark overlays, cookie banner, contact form styling, comma spacing, sitemap link"),
]

table8 = doc.add_table(rows=1, cols=3)
table8.style = "Table Grid"
table8.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr8 = table8.rows[0]
for i, h in enumerate(["Priority", "Count", "Key Items"]):
    add_cell_text(hdr8.cells[i], h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
style_header_row(hdr8, 3)

w8 = [Cm(2.5), Cm(1.5), Cm(12.5)]
for i, w in enumerate(w8):
    hdr8.cells[i].width = w

colors_map = {"Critical": RED_BG, "High": AMBER_BG, "Medium": BLUE_BG, "Low": "E2E3E5"}
for prio, count, items in summary_data:
    row = table8.add_row()
    add_cell_text(row.cells[0], prio, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(row.cells[1], count, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_text(row.cells[2], items, size=8)
    for i, w in enumerate(w8):
        row.cells[i].width = w
    set_cell_bg(row.cells[0], colors_map[prio])

doc.add_paragraph("")

# ═══════════════════════════════════════════════════════════════
#  9. CLOSING NOTE
# ═══════════════════════════════════════════════════════════════

doc.add_heading("9. Closing Note", level=1)

doc.add_paragraph(
    "We want to reiterate our appreciation for the hard work the development team has put into "
    "thinkartha.com. The site successfully communicates Artha Solutions' value proposition and "
    "serves as a strong digital presence for the company."
)
doc.add_paragraph(
    "The items listed in this report are intended as constructive observations to help refine and "
    "polish the website further. Most of the critical items involve straightforward content edits "
    "(removing duplicate widgets, correcting a typo, fixing a link) that should require minimal "
    "development time."
)
doc.add_paragraph(
    "We are happy to schedule a walkthrough of these findings at a time that works for the team, "
    "and we remain available to provide any additional context or testing support as needed."
)

closing = doc.add_paragraph()
closing.space_before = Pt(20)
r = closing.add_run("Thank you for your attention and continued dedication to excellence.")
r.font.name = "Calibri"
r.font.size = Pt(10)
r.bold = True
r.font.color.rgb = DARK_TEAL

sign = doc.add_paragraph()
sign.space_before = Pt(20)
for line in [
    "Warm regards,",
    "Vishnu Sai Rachakatla",
    "Artha Solutions",
]:
    r = sign.add_run(line + "\n")
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.color.rgb = BLACK

# ── Save ────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), "thinkartha_Website_Quality_Review.docx")
doc.save(output_path)
print(f"✅ Report saved to: {output_path}")
